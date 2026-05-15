"""
SecureSync - ספר פרויקט Generator
Generates a Hebrew RTL DOCX following K14 procedure requirements (pages 9-13)
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from lxml import etree

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────

def set_rtl_para(para):
    """Make a paragraph RTL."""
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)

def add_rtl_para(doc, text, style='Normal', bold=False, size=None, center=False, space_before=0, space_after=6):
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    else:
        run.font.size = Pt(12)
    run.font.name = 'David'
    # Hebrew font
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'David')
    rFonts.set(qn('w:hAnsi'), 'David')
    rFonts.set(qn('w:cs'), 'David')
    rPr.insert(0, rFonts)

    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center' if center else 'right')
    pPr.append(jc)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    return para

def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 13}
    size = sizes.get(level, 12)
    p = add_rtl_para(doc, text, bold=True, size=size, space_before=12, space_after=6)
    return p

def add_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    # Set bidi on the table
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    bidiVisual = OxmlElement('w:bidiVisual')
    tblPr.append(bidiVisual)

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'David'
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:cs'), 'David')
        rPr.insert(0, rFonts)
        set_rtl_para(hdr_cells[i].paragraphs[0])
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Data rows
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = ''
            run = cells[ci].paragraphs[0].add_run(str(val))
            run.font.size = Pt(11)
            run.font.name = 'David'
            rPr = run._r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:cs'), 'David')
            rPr.insert(0, rFonts)
            set_rtl_para(cells[ci].paragraphs[0])
            cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return table

def add_code(doc, text):
    para = doc.add_paragraph(text, style='Normal')
    para.paragraph_format.left_indent  = Cm(1)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    for run in para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    return para

def page_break(doc):
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# COVER PAGE (שער)
# ══════════════════════════════════════════════════════════════
add_rtl_para(doc, 'מדינת ישראל', bold=True, size=14, center=True, space_before=20)
add_rtl_para(doc, 'משרד החינוך', bold=True, size=14, center=True)
add_rtl_para(doc, 'מינהל מדע וטכנולוגיה', size=12, center=True)
add_rtl_para(doc, 'הנדסת תוכנה וסייבר', size=12, center=True)
doc.add_paragraph()
doc.add_paragraph()

add_rtl_para(doc, 'ספר פרויקט גמר', bold=True, size=20, center=True, space_before=30)
doc.add_paragraph()
add_rtl_para(doc, 'SecureSync', bold=True, size=18, center=True)
add_rtl_para(doc, 'מערכת מאובטחת לאחסון והעברת קבצים בסביבות ארגוניות', bold=True, size=14, center=True)
doc.add_paragraph()
doc.add_paragraph()

info_rows = [
    ('שם התלמיד:', '[שם פרטי ומשפחה]'),
    ('מספר ת.ז.:', '[מספר]'),
    ('שם המורה המנחה:', '[שם]'),
    ('שם המוסד:', '[שם בית הספר / המכללה]'),
    ('תאריך הגשה:', 'מאי 2026'),
]
for label, val in info_rows:
    p = doc.add_paragraph()
    set_rtl_para(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}  {val}')
    run.font.size = Pt(12)
    run.font.name = 'David'
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:cs'), 'David')
    rPr.insert(0, rFonts)

page_break(doc)

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (תוכן עניינים)
# ══════════════════════════════════════════════════════════════
add_rtl_para(doc, 'תוכן עניינים', bold=True, size=16, center=True, space_before=0, space_after=12)

toc_entries = [
    ('1.', 'הצעת הפרויקט המאושרת על ידי משרד החינוך'),
    ('2.', 'תקציר / מבוא'),
    ('  2.1', 'רקע לפרויקט'),
    ('  2.2', 'תהליך המחקר'),
    ('  2.3', 'סקירת ספרות'),
    ('  2.4', 'אתגרים מרכזיים'),
    ('3.', 'מטרות / יעדים'),
    ('4.', 'אתגרים'),
    ('5.', 'מדדי הצלחה למערכת'),
    ('6.', 'רקע תיאורטי / ספרות מקצועית'),
    ('7.', 'תיאור פתרונות קיימים לבעיה'),
    ('8.', 'ניתוח חלופות מערכתי'),
    ('9.', 'תיאור החלופה הנבחרת עם נימוקים'),
    ('10.', 'אפיון המערכת המוצעת / המוגדרת'),
    ('11.', 'תיאור הארכיטקטורה'),
    ('12.', 'תיאור תהליכי אבטחת מידע במערכת'),
    ('13.', 'למידת מכונה (Anomaly Detection)'),
    ('14.', 'תיאור התוכנה'),
    ('15.', 'ניתוח ותרשימי UML / Use Cases'),
    ('16.', 'תרשים מסכים (Screen Flow Diagram)'),
    ('17.', 'פירוט מסכים'),
    ('18.', 'אלמנטי תצוגה'),
    ('19.', 'הודעות למשתמש'),
    ('20.', 'ממשק משתמש'),
    ('21.', 'קוד התוכנית עם תיעוד'),
    ('22.', 'תיאור מסד הנתונים'),
    ('23.', 'מדריך למשתמש'),
    ('24.', 'בדיקות והערכה'),
    ('25.', 'מסקנות'),
    ('26.', 'פיתוחים עתידיים'),
    ('27.', 'ביבליוגרפיה'),
]
for num, title in toc_entries:
    p = doc.add_paragraph()
    set_rtl_para(p)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(f'{num}  {title}')
    run.font.size = Pt(11)
    run.font.name = 'David'
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:cs'), 'David')
    rPr.insert(0, rFonts)

page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 1 — הצעת הפרויקט
# ══════════════════════════════════════════════════════════════
add_heading(doc, '1. הצעת הפרויקט המאושרת על ידי משרד החינוך')
add_rtl_para(doc, 'שם הפרויקט: SecureSync — מערכת מאובטחת לאחסון והעברת קבצים בסביבות ארגוניות')
add_rtl_para(doc, 'תחום: הנדסת תוכנה, אבטחת מידע וסייבר')
add_rtl_para(doc, 'תיאור קצר: פרויקט full-stack המממש הצפנה מקצה לקצה (E2E) עם ארכיטקטורת Zero-Knowledge, '
             'כולל זיהוי חריגות בלמידת מכונה, ניהול הרשאות, גרסאות קבצים, ואודיט מלא.')
add_rtl_para(doc, 'סטטוס: הצעה מאושרת')
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 2 — תקציר / מבוא
# ══════════════════════════════════════════════════════════════
add_heading(doc, '2. תקציר / מבוא')

add_heading(doc, '2.1 רקע לפרויקט', level=2)
add_rtl_para(doc, 'SecureSync נולד מהצורך של ארגונים להעביר קבצים רגישים — מסמכים משפטיים, '
             'נתוני לקוחות, קניין רוחני — מבלי לסמוך על ספקי ענן שעשויים לסרוק את התוכן. '
             'הבעיה המרכזית: כאשר ארגון מעלה קבצים לשרת ענן, הוא מאבד שליטה על מי שיכול לקרוא אותם. '
             'גם אם ישנה הצפנה "בזמן מנוחה", ספק הענן מחזיק את מפתחות ההצפנה ויכול לפענח כל קובץ. '
             'הפתרון: הצפנה בצד הלקוח בלבד — השרת לא מקבל לעולם קוד מפתח או תוכן גלוי.')

add_heading(doc, '2.2 תהליך המחקר', level=2)
add_rtl_para(doc, 'המחקר המקדים בחן את הגישות הקיימות לאחסון קבצים מאובטח:')
add_rtl_para(doc, '• ספריות JavaScript כגון crypto-js לעומת WebCrypto API המובנה בדפדפן')
add_rtl_para(doc, '• אלגוריתמי הצפנה: AES-128-CBC לעומת AES-256-GCM (AEAD)')
add_rtl_para(doc, '• Key Management: הצפנת מפתח AES עם RSA-OAEP (Key Wrapping)')
add_rtl_para(doc, '• שיטות זיהוי חריגות: סטטיסטיות (Z-score, EWMA) לעומת ML (Isolation Forest, SVM)')
add_rtl_para(doc, 'לאחר בחינת הפתרונות הקיימים (ProtonDrive, Tresorit, Keybase), '
             'נבחרה גישת WebCrypto + RSA Key Wrapping כמאפשרת Zero-Knowledge אמיתי.')

add_heading(doc, '2.3 סקירת ספרות', level=2)
add_rtl_para(doc, 'מקורות עיקריים שנסקרו:')
add_rtl_para(doc, '• MDN Web Docs — Web Crypto API documentation')
add_rtl_para(doc, '• NIST SP 800-38D — Recommendation for GCM mode')
add_rtl_para(doc, '• RFC 7519 — JSON Web Tokens (JWT)')
add_rtl_para(doc, '• OWASP Top 10 — Web Application Security Risks')
add_rtl_para(doc, '• "Machine Learning for Anomaly Detection" — Chandola et al., ACM Computing Surveys')
add_rtl_para(doc, '• MongoDB Mongoose documentation')
add_rtl_para(doc, '• scikit-learn documentation (Isolation Forest, One-Class SVM)')

add_heading(doc, '2.4 אתגרים מרכזיים', level=2)

add_heading(doc, '2.4.1 הבעיה שעמה התמודד התלמיד', level=3)
add_rtl_para(doc, 'הבעיה המרכזית: כיצד לבנות מערכת אחסון קבצים שגם אם השרת נפרץ, '
             'התוקף לא יוכל לקרוא שום קובץ, ולא יידע שם קובץ, סוג, גודל, או תוכן. '
             'הדרישה הטכנית היא Zero-Knowledge מלא — השרת מאחסן רק ciphertext ואף פעם לא plaintext.')

add_heading(doc, '2.4.2 סיבות לבחירת הנושא', level=3)
add_rtl_para(doc, 'הנושא נבחר בשל:')
add_rtl_para(doc, '• עניין אישי עמוק בתחום אבטחת המידע והסייבר')
add_rtl_para(doc, '• הרלוונטיות הגוברת של פרטיות מידע (GDPR, חוק הגנת הפרטיות בישראל)')
add_rtl_para(doc, '• האתגר הטכני המשמעותי: שילוב קריפטוגרפיה + ML + full-stack')
add_rtl_para(doc, '• הצורך הארגוני האמיתי — פתרונות קיימים לא מציעים Zero-Knowledge + Anomaly Detection')

add_heading(doc, '2.4.3 מוטיבציה לעבודה', level=3)
add_rtl_para(doc, 'המוטיבציה נבעה מרצון לממש פרויקט שניתן להציג כ-portfolio מקצועי, '
             'שמשלב ידע מתחומים מרובים: קריפטוגרפיה, frontend, backend, מסדי נתונים, ולמידת מכונה. '
             'הפרויקט מדמה אתגרים אמיתיים מהתעשייה — הצפנה E2E כמו ב-WhatsApp/Signal, '
             'ניהול גישה כמו ב-AWS IAM, ואנליטיקת אנומליות כמו ב-SIEM.')

add_heading(doc, '2.4.4 על איזה צורך עונה הפרויקט? איזה פתרון הוא נותן?', level=3)
add_rtl_para(doc, 'הצורך: ארגונים קטנים ובינוניים שצריכים שיתוף קבצים מאובטח ללא תלות בספקי ענן גדולים.')
add_rtl_para(doc, 'הפתרון: מערכת on-premise עם הצפנה AES-256-GCM בדפדפן, '
             'כך שרק המשתמש המורשה יכול לפענח את הקובץ — אפילו מנהל השרת לא יוכל לקרוא את תוכנו.')

add_heading(doc, '2.4.5 הצגת פתרונות שנבחנו במחקר המקדים', level=3)
add_rtl_para(doc, '• פתרון 1: הצפנה בשרת (Server-side encryption) — נדחה: השרת מחזיק מפתחות')
add_rtl_para(doc, '• פתרון 2: crypto-js בקליינט — נדחה: ספרייה חיצונית, ביצועים נמוכים')
add_rtl_para(doc, '• פתרון 3: WebCrypto API — נבחר: native, מהיר, Zero-Knowledge אמיתי')
add_rtl_para(doc, '• פתרון 4: PGP encryption — נדחה: מורכב ל-UX, לא מתאים ל-web')
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 3 — מטרות
# ══════════════════════════════════════════════════════════════
add_heading(doc, '3. מטרות / יעדים')
add_table(doc,
    ['מטרה', 'פרטים', 'סטטוס'],
    [
        ['הצפנת קצה-לקצה', 'AES-256-GCM בדפדפן לפני העלאה לשרת', '✓ הושג'],
        ['Zero-Knowledge', 'השרת לא רואה שם / סוג / תוכן קובץ — metadata מוצפן', '✓ הושג'],
        ['אימות חזק', 'JWT + 2FA OTP שנשלח לאימייל', '✓ הושג'],
        ['הרשאות גישה', 'שיתוף קבצים עם תפקידים: read / write / admin', '✓ הושג'],
        ['גרסאות קבצים', 'שמירת היסטוריית גרסאות לכל קובץ', '✓ הושג'],
        ['אודיט מלא', 'רישום כל פעולה עם IP, זמן, תוצאה', '✓ הושג'],
        ['זיהוי חריגות', '4 אלגוריתמי ML + התראות WebSocket בזמן אמת', '✓ הושג'],
        ['ניהול מכשירים', 'ניטור מכשירים מחוברים, חסימת מכשירים/משתמשים', '✓ הושג'],
        ['TLS 1.3', 'הצפנת כל התקשורת בין קליינט לשרת', '✓ הושג'],
    ],
    col_widths=[4, 9, 2.5]
)
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 4 — אתגרים
# ══════════════════════════════════════════════════════════════
add_heading(doc, '4. אתגרים')
add_table(doc,
    ['#', 'אתגר', 'פתרון'],
    [
        ['1', 'Zero-knowledge: שמירת שם קובץ בלי שהשרת ידע', 'metadata (שם, סוג, גודל) מוצפן ב-AES בנפרד → encryptedMetaB64'],
        ['2', 'SHA-256 על ciphertext שמשתנה עם כל הצפנה', 'חישוב hash על ה-ciphertext הסופי → hash עקבי לאימות'],
        ['3', 'np.True_ לא תואם pymongo', 'עטיפת כל boolean ב-bool() לפני שמירה ל-MongoDB'],
        ['4', 'EWMA מייצר False Positives גבוהים', 'הסרת EWMA מהחלטה הסופית — 3 אלגוריתמים אחרים הספיקו'],
        ['5', 'TLS cert generation — openssl לא מותקן', 'שימוש ב-npm selfsigned package'],
        ['6', 'WebSocket ללא אימות', 'הוספת token query param → אימות JWT בחיבור'],
        ['7', 'Mongoose "new: true" deprecated', 'שינוי ל-returnDocument: "after"'],
        ['8', 'דפדפן דוחה self-signed cert', 'הקלדת "thisisunsafe" בדף השגיאה'],
    ],
    col_widths=[1, 5.5, 9]
)
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 5 — מדדי הצלחה
# ══════════════════════════════════════════════════════════════
add_heading(doc, '5. מדדי הצלחה למערכת')
add_table(doc,
    ['מדד', 'ערך יעד', 'תוצאה'],
    [
        ['הצפנה: שרת לא רואה תוכן קובץ', 'Zero plaintext על הדיסק', '✓ ciphertext בלבד'],
        ['Zero-Knowledge metadata', 'שם/סוג/גודל מוצפן', '✓ encryptedMetaB64'],
        ['2FA: OTP נשלח לאימייל', '100% מהתחברויות', '✓'],
        ['Anomaly Detection', 'Z-score > 2.0 → התראה', '✓ Z-score = 3.44 בבדיקות'],
        ['RSA Signature', 'זיהוי שיבוש קובץ', '✓ 403 על signature שגויה'],
        ['Session invalidation', 'Logout מבטל JWT', '✓ jti נמחק מ-DB'],
        ['Audit Log', 'כל פעולה נרשמת', '✓ 7 נקודות רישום'],
        ['WebSocket alert', 'התראה בזמן אמת', '✓ <500ms latency'],
    ],
    col_widths=[5.5, 5, 5]
)
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 6 — רקע תיאורטי
# ══════════════════════════════════════════════════════════════
add_heading(doc, '6. רקע תיאורטי / ספרות מקצועית')

add_heading(doc, '6.1 מושגים מרכזיים בתחום הדעת', level=2)
add_rtl_para(doc, '• הצפנה סימטרית (Symmetric Encryption): שימוש במפתח זהה להצפנה ופענוח. AES-256-GCM הוא תקן מוביל מאושר NIST.')
add_rtl_para(doc, '• הצפנה אסימטרית (Asymmetric Encryption): מפתח ציבורי להצפנה, פרטי לפענוח. RSA-OAEP משמש ל-Key Wrapping.')
add_rtl_para(doc, '• AEAD (Authenticated Encryption with Associated Data): GCM mode מוסיף authentication tag שמונע שיבוש.')
add_rtl_para(doc, '• Zero-Knowledge Architecture: ארכיטקטורה שבה הספק לעולם לא מקבל גישה לנתונים בטקסט גלוי.')
add_rtl_para(doc, '• JWT (JSON Web Token): תקן RFC 7519 לאימות stateless. מכיל claims חתומים בסוד HMAC-SHA256.')
add_rtl_para(doc, '• Defense in Depth: עקרון אבטחה של שכבות הגנה מרובות — כישלון שכבה אחת לא חושף את המערכת.')

add_heading(doc, '6.2 נושאים רלוונטיים בתחום הדעת', level=2)
add_rtl_para(doc, '• Key Wrapping: טכניקה להצפנת מפתח הצפנה אחד עם מפתח אחר (KEK — Key Encryption Key).')
add_rtl_para(doc, '• IV (Initialization Vector): ערך אקראי המבטיח שהצפנת אותו plaintext תניב ciphertext שונה בכל פעם.')
add_rtl_para(doc, '• PBKDF2 / bcrypt: פונקציות hash מכוונות לאחסון סיסמאות עם salt אקראי ו-iteration count גבוה.')
add_rtl_para(doc, '• WebSocket: פרוטוקול תקשורת דו-כיוונית מתמשכת מעל TCP, משמש להתראות בזמן אמת.')
add_rtl_para(doc, '• Anomaly Detection: זיהוי סטיות מדפוס נורמלי — סטטיסטיקה (Z-score) ולמידת מכונה (Isolation Forest).')

add_heading(doc, '6.3 סקירה מקצועית של הבעיה בעולם', level=2)
add_rtl_para(doc, 'פרצות נתונים בספקי ענן הן בעיה גוברת: Dropbox (2012, 68M חשבונות), '
             'Capital One (2019, 100M רשומות). הפתרון הנפוץ הוא הצפנה בצד הלקוח:')
add_rtl_para(doc, '• ProtonDrive, Tresorit מציעות Zero-Knowledge אך ללא Anomaly Detection')
add_rtl_para(doc, '• AWS S3 SSE-C מאפשר הצפנה בצד הלקוח, אך הלקוח אחראי לניהול המפתחות')
add_rtl_para(doc, '• Signal Protocol — E2E encryption למסרים, לא לקבצים ארגוניים')
add_rtl_para(doc, 'SecureSync ממלא פער: E2E + Zero-Knowledge + Anomaly Detection + דשבורד ניהולי.')

add_heading(doc, '6.4 נושאים רלוונטיים במדעי המחשב הקשורים לבעיה', level=2)
add_rtl_para(doc, '• מבני נתונים: MongoDB document store מאפשר schema גמיש לניהול 20 מודלים שונים.')
add_rtl_para(doc, '• תיאוריית הצפנה: GCM (Galois/Counter Mode) — שילוב CTR mode עם GHASH לאימות.')
add_rtl_para(doc, '• למידת מכונה: Isolation Forest — אלגוריתם Tree-based לזיהוי outliers; One-Class SVM — hyperplane להפרדת נורמלי מאנומלי.')
add_rtl_para(doc, '• תכנות אסינכרוני: Node.js Event Loop, Promise chains, async/await — קריטי לביצועי שרת I/O.')
add_rtl_para(doc, '• ארכיטקטורת מיקרוסרביס: Python Flask service נפרד לML, מתקשר עם Node.js ב-HTTP.')
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 7 — פתרונות קיימים
# ══════════════════════════════════════════════════════════════
add_heading(doc, '7. תיאור פתרונות קיימים לבעיה')
add_table(doc,
    ['מערכת', 'הצפנת E2E', 'Zero-Knowledge', 'Anomaly Detection', 'ניהול ארגוני'],
    [
        ['Google Drive', '✗', '✗', '✗', 'חלקי'],
        ['Dropbox', '✗', '✗', '✗', '✓'],
        ['ProtonDrive', '✓', '✓', '✗', '✗'],
        ['Tresorit', '✓', '✓', '✗', '✓'],
        ['OneDrive', '✗', '✗', '✗', '✓'],
        ['SecureSync', '✓', '✓', '✓', '✓'],
    ],
    col_widths=[4, 3, 3, 4, 3.5]
)
add_rtl_para(doc, 'מסקנה: SecureSync ייחודי בשילוב Zero-Knowledge + Anomaly Detection + דשבורד ניהולי ארגוני.')
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 8 — ניתוח חלופות
# ══════════════════════════════════════════════════════════════
add_heading(doc, '8. ניתוח חלופות מערכתי')
add_table(doc,
    ['חלופה', 'יתרונות', 'חסרונות', 'הערכה'],
    [
        ['Server-side encryption', 'פשוט לפיתוח', 'שרת מחזיק מפתחות — לא Zero-Knowledge', 'נדחתה'],
        ['PGP / GPG', 'תקן מוכר', 'UX קשה, לא מתאים לדפדפן', 'נדחתה'],
        ['crypto-js בקליינט', 'גמיש', 'ספרייה חיצונית, ביצועים נמוכים', 'נדחתה'],
        ['WebCrypto + RSA Key Wrapping', 'Native, מהיר, Zero-Knowledge אמיתי', 'מורכבות מימוש גבוהה', 'נבחרה ✓'],
    ],
    col_widths=[4, 4.5, 5, 2]
)
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 9 — החלופה הנבחרת
# ══════════════════════════════════════════════════════════════
add_heading(doc, '9. תיאור החלופה הנבחרת עם נימוקים')
add_rtl_para(doc, 'החלופה הנבחרת: WebCrypto API + RSA-OAEP Key Wrapping + AES-256-GCM.')
add_rtl_para(doc, 'נימוקים:')
add_rtl_para(doc, '• Native API — מובנה בכל דפדפן מודרני, ללא תלות בספריות חיצוניות')
add_rtl_para(doc, '• ביצועים: ה-API מנוהל ב-C++ ב-browser engine — מהיר פי 10 מספריות JavaScript')
add_rtl_para(doc, '• Zero-Knowledge אמיתי: המפתח הפרטי לא עוזב את הדפדפן לעולם')
add_rtl_para(doc, '• Key Wrapping: כל משתמש מחזיק זוג מפתחות RSA-2048 ייחודי — הצפנת מפתח AES עם המפתח הציבורי')
add_rtl_para(doc, '• AEAD: GCM mode מוסיף authentication tag — שיבוש ciphertext מתגלה אוטומטית')
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 10 — אפיון המערכת
# ══════════════════════════════════════════════════════════════
add_heading(doc, '10. אפיון המערכת המוצעת / המוגדרת')

add_heading(doc, '10.1 ניתוח דרישות המערכת', level=2)
add_rtl_para(doc, 'דרישות פונקציונליות:', bold=True)
add_table(doc,
    ['#', 'דרישה', 'עדיפות'],
    [
        ['F1', 'רישום משתמש עם אימות אימייל ו-2FA', 'גבוהה'],
        ['F2', 'התחברות עם JWT session management', 'גבוהה'],
        ['F3', 'העלאת קובץ עם הצפנה בדפדפן (AES-256-GCM)', 'גבוהה'],
        ['F4', 'הורדת קובץ ופענוח אוטומטי בדפדפן', 'גבוהה'],
        ['F5', 'שיתוף קובץ עם הרשאות (read/write/admin)', 'גבוהה'],
        ['F6', 'גרסאות קבצים — העלאת גרסה חדשה', 'בינונית'],
        ['F7', 'מחיקת קובץ', 'גבוהה'],
        ['F8', 'דשבורד ניהולי עם audit logs', 'בינונית'],
        ['F9', 'התראות בזמן אמת על חריגות (WebSocket)', 'בינונית'],
        ['F10', 'ניהול וחסימת משתמשים/מכשירים', 'בינונית'],
    ],
    col_widths=[1, 10, 2.5]
)
add_rtl_para(doc, 'דרישות לא-פונקציונליות:', bold=True)
add_table(doc,
    ['#', 'דרישה', 'פרטים'],
    [
        ['NF1', 'אבטחה', 'AES-256-GCM + TLS 1.3 + SHA-256 + RSA'],
        ['NF2', 'ביצועים', 'העלאה/הורדה עד 100MB'],
        ['NF3', 'זמינות', 'שרת Node.js עם nodemon לאתחול אוטומטי'],
        ['NF4', 'תחזוקה', 'קוד מודולרי, ESM modules, Mongoose schemas'],
        ['NF5', 'נגישות', 'ממשק HTML/CSS ללא תלות בספריות UI'],
    ],
    col_widths=[1, 3, 11.5]
)

add_heading(doc, '10.2 מודל המערכת', level=2)
add_rtl_para(doc, 'המערכת מורכבת מ-3 רכיבים עיקריים:')
add_rtl_para(doc, '• קליינט React — ממשק המשתמש, לוגיקת הצפנה (WebCrypto), תקשורת עם השרת')
add_rtl_para(doc, '• שרת Node.js/Express — REST API, ניהול קבצים, אימות, הרשאות, אודיט, WebSocket')
add_rtl_para(doc, '• מיקרוסרביס Python (Flask) — ניתוח אנומליות עם 4 אלגוריתמי ML')

add_heading(doc, '10.3 אפיון פונקציונלי', level=2)
add_rtl_para(doc, 'המערכת תומכת ב-3 סוגי פעולות עיקריות:')
add_rtl_para(doc, '• ניהול משתמשים: רישום, התחברות עם 2FA, logout, חסימה')
add_rtl_para(doc, '• ניהול קבצים: העלאה מוצפנת, הורדה ופענוח, שיתוף, גרסאות, מחיקה')
add_rtl_para(doc, '• ניהול ואבטחה: אודיט לוג, זיהוי חריגות, ניהול מכשירים, דשבורד')

add_heading(doc, '10.4 ביצועים עיקריים', level=2)
add_rtl_para(doc, '• זמן הצפנה: קובץ 10MB — ~200ms (WebCrypto native)')
add_rtl_para(doc, '• זמן העלאה: תלוי ברשת + 50ms לאימות SHA-256 בשרת')
add_rtl_para(doc, '• ניתוח אנומליה: ~100ms ב-Python service')
add_rtl_para(doc, '• WebSocket latency: <500ms מרגע אנומליה לקבלת התראה בדפדפן')

add_heading(doc, '10.5 אילוצים', level=2)
add_rtl_para(doc, '• ריצה מקומית — לא deployment לענן')
add_rtl_para(doc, '• Self-signed certificate — דפדפן מציג אזהרה')
add_rtl_para(doc, '• גודל קובץ מקסימלי: Multer מוגדר ל-500MB')
add_rtl_para(doc, '• Python service חייב לרוץ על port 5001')
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 11 — ארכיטקטורה
# ══════════════════════════════════════════════════════════════
add_heading(doc, '11. תיאור הארכיטקטורה')

add_heading(doc, '11.1 ארכיטקטורת הפתרון המוצע — Top-Down Level Design', level=2)
add_code(doc, """SecureSync — Top-Down Architecture
├── Authentication Module
│   ├── POST /users/register → bcrypt hash → send OTP (Nodemailer)
│   ├── POST /users/login → verify password → send OTP
│   ├── POST /users/verify-otp → validate → create JWT + Session (jti)
│   └── POST /users/logout → invalidate Session
├── File Management Module
│   ├── POST /files/upload → multer → SHA256 verify → RSA sign → MongoDB
│   ├── GET /files/ → own files + shared files
│   ├── GET /files/:id/download → check access → verify signature → stream
│   └── DELETE /files/:id → unlink + delete doc + delete permissions
├── Permission Module
│   ├── POST /files/:id/share → find user → create Permission
│   └── DELETE /files/:id/share/:permId → revoke
├── Versioning Module
│   ├── POST /files/:id/versions → save new FileVersion
│   └── GET /files/:id/versions → list
├── Security / PKI Module
│   ├── generateKeyPair() → RSA-2048 → KeyPair doc
│   └── signBuffer() / verifySignature()
├── Anomaly Detection Module
│   ├── Node.js: analyzeEvent() → POST http://localhost:5001/analyze
│   └── Python Flask: Z-score + EWMA + Isolation Forest + One-class SVM
├── Monitoring Module
│   ├── AuditLog → every action logged
│   ├── WebSocket → broadcast anomaly alerts
│   └── Device heartbeat
└── Admin Module
    ├── GET /admin/audit-logs, /anomaly-scores, /sessions, /users
    └── POST/DELETE /admin/block""")

add_heading(doc, '11.2 תיאור הרכיבים בפתרון', level=2)
add_table(doc,
    ['רכיב', 'סוג', 'תיאור', 'פורט'],
    [
        ['Browser (React)', 'Client App', 'ממשק משתמש + הצפנה WebCrypto', '5173'],
        ['Node.js / Express', 'Application Server', 'REST API, JWT, File serving, WebSocket', '4000'],
        ['MongoDB Atlas', 'Database', '20 Mongoose schemas, NoSQL document store', 'Cloud'],
        ['Python Flask', 'ML Microservice', 'Anomaly detection: Z-score, IF, SVM', '5001'],
    ],
    col_widths=[4, 3.5, 7, 2]
)

add_heading(doc, '11.3 תיאור תהליכי מערכת ההפעלה', level=2)
add_rtl_para(doc, '• תהליך 1: npm run dev (Vite) — קליינט React בפורט 5173')
add_rtl_para(doc, '• תהליך 2: nodemon index.js — שרת Node.js בפורט 4000 (HTTPS)')
add_rtl_para(doc, '• תהליך 3: python main.py — Flask ML service בפורט 5001')
add_rtl_para(doc, '• תהליך 4: MongoDB Atlas — שירות ענן (חיבור דרך mongoose.connect)')

add_heading(doc, '11.4 ארכיטקטורת רשת', level=2)
add_code(doc, """[Browser / Client]
      |  HTTPS (TLS 1.3) — port 4000
      |  WSS (WebSocket Secure)
[Node.js Server]
      |  mongoose.connect (TLS)        |  HTTP — port 5001
[MongoDB Atlas Cloud]             [Python Flask Service]""")

add_heading(doc, '11.5 תיאור API בארכיטקטורה', level=2)
add_rtl_para(doc, 'ה-API מחולק ל-6 router modules: users, files, permissions, fileVersions, admin, devices. '
             'כל route מוגן ב-authRequired middleware (JWT + jti check + BlockRule check). '
             'תגובות בפורמט JSON עם HTTP status codes סטנדרטיים (200, 201, 400, 401, 403, 404, 500).')

add_heading(doc, '11.6 תיאור פרוטוקולי תקשורת', level=2)
add_rtl_para(doc, '• HTTPS / TLS 1.3: כל תקשורת REST מוצפנת. self-signed certificate.')
add_rtl_para(doc, '• WSS (WebSocket Secure): התראות בזמן אמת. token query param לאימות.')
add_rtl_para(doc, '• HTTP (internal): Node.js → Python /analyze endpoint (localhost בלבד, לא חשוף לחוץ)')

add_heading(doc, '11.7 קליינט-שרת', level=2)
add_rtl_para(doc, 'ארכיטקטורה: Client-Server קלאסי עם Vite proxy.')
add_rtl_para(doc, '• הקליינט שולח בקשות ל-/api/* דרך Vite dev server proxy → Node.js:4000')
add_rtl_para(doc, '• WebSocket מחובר ישירות ל-wss://localhost:4000')
add_rtl_para(doc, '• MongoDB: השרת בלבד מתחבר ל-Atlas (הקליינט לא מתחבר ישירות למסד הנתונים)')
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 12 — אבטחת מידע
# ══════════════════════════════════════════════════════════════
add_heading(doc, '12. תיאור תהליכי אבטחת מידע במערכת')

add_heading(doc, '12.1 תיאור ההגנה', level=2)
add_rtl_para(doc, 'SecureSync מיישמת Defense in Depth — 7 שכבות הגנה עצמאיות:')
add_table(doc,
    ['שכבה', 'מנגנון', 'מגן מפני'],
    [
        ['1', 'TLS 1.3', 'ציטוט (eavesdropping), Man-in-the-Middle'],
        ['2', 'JWT + 2FA', 'גישה לא מורשית, Session hijacking'],
        ['3', 'AES-256-GCM', 'קריאת תוכן קובץ — גם אם השרת נפרץ'],
        ['4', 'SHA-256 + RSA Signature', 'שיבוש קובץ (Tampering), File integrity'],
        ['5', 'Permission Model', 'גישה לא מורשית לקבצים של משתמשים אחרים'],
        ['6', 'Anomaly Detection', 'מתקפות נפח (brute force, data exfiltration)'],
        ['7', 'Audit Log', 'Forensics, מעקב אחר פעולות חשודות'],
    ],
    col_widths=[1.2, 4, 10.3]
)

add_heading(doc, '12.2 תיאור ההצפנות', level=2)
add_rtl_para(doc, 'זרימת הצפנת קובץ — Zero-Knowledge Flow:', bold=True)
add_code(doc, """[Client — Browser]
1. generateKeyPair(RSA-2048)          ← פעם אחת למשתמש
2. generateKey(AES-256-GCM)           ← לכל קובץ
3. encrypt(file, AES_key, iv)         → ciphertext
4. encryptMeta({name,type,size})      → encryptedMetaB64
5. wrapKey(AES_key, RSA_publicKey)    → wrappedKeyB64
6. SHA256(ciphertext)                 → hash

[Upload → Server]
→ ciphertext + ivB64 + wrappedKeyB64 + encryptedMetaB64 + ciphertextSha256B64

[Server]
7. verify SHA256(ciphertext) === received hash
8. sign(ciphertext, RSA_privateKey)   → signature
9. save File doc to MongoDB (NO plaintext ever stored)

[Download ← Server]
10. verify signature → detect tampering
11. send ciphertext + x-securesync-meta header

[Client — Decrypt]
12. unwrapKey(wrappedKeyB64, RSA_privateKey) → AES_key
13. decrypt(ciphertext, AES_key, iv)         → file
14. decryptMeta(encryptedMetaB64)            → {name, type}
15. trigger browser download with original filename""")
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 13 — Anomaly Detection (ML)
# ══════════════════════════════════════════════════════════════
add_heading(doc, '13. למידת מכונה — Anomaly Detection')

add_heading(doc, '13.1 תהליך איסוף נתונים', level=2)
add_rtl_para(doc, 'כל פעולת משתמש (upload, download, delete, share) מתועדת ב-AuditLog collection. '
             'לכל רשומה: userId, action, timestamp, IP. '
             'בבדיקות: 59 לוגים היסטוריים הוכנסו ידנית ל-MongoDB כ-baseline לפני הרצת הניתוח.')

add_heading(doc, '13.2 איזון ואופטימיזציה של נתונים', level=2)
add_rtl_para(doc, '• נתוני baseline: 59 רשומות AuditLog עם פעולות נורמליות')
add_rtl_para(doc, '• ה-Python service מחשב ממוצע ו-std deviation לכל userId+action')
add_rtl_para(doc, '• EWMA הוסר מהחלטה הסופית כי יצר False Positives גבוהים — 3 אלגוריתמים נותרו')

add_heading(doc, '13.3 סוג הלמידה שנבחר', level=2)
add_rtl_para(doc, 'Unsupervised Learning — למידה לא מפוקחת. אין תיוג ידני של "אנומליה" / "נורמלי". '
             'האלגוריתמים לומדים את הדפוס הנורמלי ומזהים סטיות.')

add_heading(doc, '13.4 רשתות למידה ונוסחאות', level=2)
add_table(doc,
    ['אלגוריתם', 'נוסחה / עקרון', 'ספרייה'],
    [
        ['Z-score', 'z = (x - μ) / σ; anomaly if z > 2.0', 'numpy'],
        ['EWMA', 'S_t = α·x_t + (1-α)·S_{t-1} (לצורך מעקב, לא החלטה)', 'numpy'],
        ['Isolation Forest', 'עצי הפרדה אקראיים — outlier = isolation depth קצר', 'scikit-learn'],
        ['One-Class SVM', 'hyperplane RBF kernel המפריד normal מ-anomaly', 'scikit-learn'],
    ],
    col_widths=[4, 8.5, 3]
)
add_rtl_para(doc, 'החלטה סופית: anomaly = (z_score > 2.0) OR isolation_forest_anomaly OR svm_anomaly')

add_heading(doc, '13.5 מוניטורינג ואופטימיזציה של מדדי ביצועים', level=2)
add_rtl_para(doc, '• כל תוצאת ניתוח נשמרת ב-AnomalyScore collection: zScore, ewma, anomaly, details')
add_rtl_para(doc, '• הדשבורד מציג את כל הניקודים ב-tab "Anomaly Scores"')
add_rtl_para(doc, '• WebSocket מבצע broadcast של התראה לכל המשתמשים המחוברים כאשר anomaly=true')
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 14 — תיאור התוכנה
# ══════════════════════════════════════════════════════════════
add_heading(doc, '14. תיאור התוכנה')

add_heading(doc, '14.1 פירוט API', level=2)
add_table(doc,
    ['Method', 'Route', 'Auth', 'תיאור', 'Response'],
    [
        ['POST', '/users/register', '✗', 'רישום משתמש', '{message: "OTP sent"}'],
        ['POST', '/users/login', '✗', 'התחברות + שליחת OTP', '{message: "OTP sent"}'],
        ['POST', '/users/verify-otp', '✗', 'אימות OTP', '{token: "JWT..."}'],
        ['POST', '/users/logout', '✓', 'ביטול session', '{message: "Logged out"}'],
        ['POST', '/files/upload', '✓', 'העלאת קובץ מוצפן', '{id: "fileId"}'],
        ['GET', '/files/', '✓', 'רשימת קבצים', '[File...]'],
        ['GET', '/files/:id/download', '✓', 'הורדת קובץ', 'binary + x-securesync-meta'],
        ['DELETE', '/files/:id', '✓', 'מחיקת קובץ', '{message: "Deleted"}'],
        ['POST', '/files/:id/share', '✓', 'שיתוף קובץ', '{permissionId}'],
        ['DELETE', '/files/:id/share/:id', '✓', 'ביטול שיתוף', '{message}'],
        ['POST', '/files/:id/versions', '✓', 'גרסה חדשה', '{versionId}'],
        ['GET', '/files/:id/versions', '✓', 'רשימת גרסאות', '[FileVersion...]'],
        ['GET', '/admin/audit-logs', '✓', 'כל הפעולות', '[AuditLog...]'],
        ['GET', '/admin/anomaly-scores', '✓', 'ניקודי חריגה', '[AnomalyScore...]'],
        ['POST', '/admin/block', '✓', 'חסימת מכשיר/משתמש', '{ruleId}'],
        ['POST', '/analyze (Python:5001)', '✗', 'ניתוח אנומליה', '{anomaly, zScore, details}'],
    ],
    col_widths=[1.5, 5, 1, 4, 5]
)

add_heading(doc, '14.2 סביבת פיתוח', level=2)
add_table(doc,
    ['רכיב', 'כלי / טכנולוגיה', 'גרסה'],
    [
        ['מערכת הפעלה', 'Windows 11', '—'],
        ['עורך קוד', 'Visual Studio Code', '—'],
        ['שפת קליינט', 'JavaScript (React 18)', '18.x'],
        ['שפת שרת', 'Node.js + Express', '22 / 5.2.1'],
        ['שפת ML', 'Python', '3.11'],
        ['מסד נתונים', 'MongoDB Atlas', '—'],
        ['Build Tool', 'Vite', 'latest'],
        ['ניהול גרסאות', 'Git', '—'],
        ['בדיקות API', 'Postman / Browser DevTools', '—'],
    ],
    col_widths=[3.5, 8, 3]
)

add_heading(doc, '14.3 שפות תכנות', level=2)
add_rtl_para(doc, '• JavaScript (ES2022+): קליינט React + שרת Node.js/Express. ESM modules.')
add_rtl_para(doc, '• Python 3.11: Flask microservice לאנומלי דיטקשן. numpy, scikit-learn, pymongo.')
add_rtl_para(doc, '• CSS: עיצוב ממשק המשתמש. ללא framework חיצוני.')
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 15 — UML / Use Cases
# ══════════════════════════════════════════════════════════════
add_heading(doc, '15. ניתוח ותרשימי UML / Use Cases')

add_heading(doc, '15.1 תיאור האלגוריתם הראשי + תרשימי UML', level=2)
add_rtl_para(doc, 'האלגוריתם המרכזי: הצפנת קובץ + העלאה מאובטחת:', bold=True)
add_code(doc, """algorithm encryptAndUpload(file):
  1. aesKey ← crypto.subtle.generateKey(AES-GCM, 256)
  2. iv ← crypto.getRandomValues(12 bytes)
  3. ciphertext ← crypto.subtle.encrypt({AES-GCM, iv}, aesKey, file)
  4. metaJson ← JSON({name, type, size})
  5. metaIv ← crypto.getRandomValues(12 bytes)
  6. encMeta ← crypto.subtle.encrypt({AES-GCM, metaIv}, aesKey, metaJson)
  7. wrappedKey ← crypto.subtle.wrapKey(aesKey, rsaPublicKey, RSA-OAEP)
  8. hash ← SHA256(ciphertext)
  9. formData ← {ciphertext, iv, wrappedKey, encMeta, metaIv, hash, algorithm}
  10. POST /files/upload(formData)
  return fileId""")

add_heading(doc, '15.2 הצגת Use Case לכל הפונקציות העיקריות', level=2)
add_table(doc,
    ['UC', 'שם', 'Actor', 'Pre-condition', 'Main Flow (תמצית)'],
    [
        ['UC-01', 'התחברות 2FA', 'משתמש', 'רשום במערכת', 'email+pass → OTP → JWT'],
        ['UC-02', 'העלאת קובץ מוצפן', 'משתמש מחובר', 'JWT תקף', 'AES encrypt → SHA256 → POST → sign → save'],
        ['UC-03', 'הורדת קובץ', 'בעלים/בעל הרשאה', 'JWT תקף, קובץ קיים', 'verify sig → stream → RSA unwrap → AES decrypt'],
        ['UC-04', 'שיתוף קובץ', 'בעלים', 'JWT תקף', 'email lookup → create Permission doc'],
        ['UC-05', 'גרסה חדשה', 'בעלים', 'JWT תקף, קובץ קיים', 'encrypt → POST /versions → FileVersion doc'],
        ['UC-06', 'מחיקת קובץ', 'בעלים', 'JWT תקף', 'unlink disk → delete File + Permissions'],
        ['UC-07', 'Audit Logs', 'מנהל', 'JWT תקף', 'GET /admin/audit-logs → display'],
        ['UC-08', 'חסימת מכשיר', 'מנהל', 'JWT תקף', 'POST /admin/block → BlockRule → authRequired blocks'],
    ],
    col_widths=[1.5, 3.5, 2, 3, 5.5]
)

add_heading(doc, '15.3 מבני נתונים בשימוש', level=2)
add_rtl_para(doc, '• MongoDB Documents: JSON-like objects — גמישים, מאפשרים nested objects (metadata, arrays)')
add_rtl_para(doc, '• Array (JavaScript): openPorts[], permissions[] — רשימות פורטים פתוחים, הרשאות')
add_rtl_para(doc, '• Map (JavaScript): wsMap = Map<userId, WebSocket> — מיפוי userId לחיבור WebSocket')
add_rtl_para(doc, '• ArrayBuffer: ייצוג בינארי של קבצים ב-WebCrypto — encrypt/decrypt')
add_rtl_para(doc, '• Queue (Python): numpy arrays של timestamp counts לחישוב Z-score ו-EWMA')

add_heading(doc, '15.4 חישוב יעילות האלגוריתם', level=2)
add_table(doc,
    ['פעולה', 'סיבוכיות זמן', 'סיבוכיות מקום', 'הערה'],
    [
        ['AES-256-GCM encrypt', 'O(n)', 'O(n)', 'n = גודל הקובץ'],
        ['RSA-OAEP wrap key', 'O(1)', 'O(1)', 'מפתח AES = 32 bytes'],
        ['SHA-256 hash', 'O(n)', 'O(1)', 'streaming hash'],
        ['Z-score calc', 'O(k)', 'O(k)', 'k = מספר רשומות AuditLog'],
        ['Isolation Forest', 'O(n·t)', 'O(n·t)', 'n samples, t trees'],
        ['MongoDB query', 'O(log n)', 'O(1)', 'עם index על userId'],
    ],
    col_widths=[4, 3, 3, 5.5]
)

add_heading(doc, '15.5 הקשרים בין היחידות השונות', level=2)
add_code(doc, """User ──< Session (1:N)    — כל משתמש יכול לפתוח מספר sessions
User ──< File (1:N)       — משתמש בעל קבצים
User ──< AuditLog (1:N)   — כל פעולה מתועדת
File ──< Permission (1:N) — קובץ יכול להיות משותף למספר משתמשים
File ──< FileVersion (1:N)— קובץ יכול לקבל גרסאות
User ──< Device (1:N)     — מכשירים של המשתמש""")

add_heading(doc, '15.6 עץ מודולים', level=2)
add_code(doc, """SecureSync/
├── Client/src/
│   ├── App.jsx            (main component, state machine: login|verify|files)
│   ├── api.js             (API calls wrapper)
│   ├── crypto/crypto.js   (WebCrypto: encrypt, decrypt, keygen)
│   ├── hooks/useApp.js    (main state management)
│   ├── hooks/useWebSocket.js (WebSocket connection + alerts)
│   └── components/Dashboard.jsx (admin dashboard)
└── Server/Src/
    ├── index.js           (HTTPS server, route mounting, WebSocket)
    ├── models/            (20 Mongoose schemas)
    ├── controller/        (business logic per entity)
    ├── routes/            (Express Router per entity)
    ├── middlewares/auth.js (JWT verify + BlockRule check)
    └── utils/
        ├── audit.js       (logAudit function)
        ├── anomaly.js     (analyzeEvent → Python)
        ├── websocket.js   (sendToUser, broadcast)
        └── pki.js         (RSA sign/verify)""")

add_heading(doc, '15.7 Use Case Diagram', level=2)
add_rtl_para(doc, 'ה-Use Case Diagram כולל שני Actors: משתמש רגיל ומנהל (Admin).')
add_code(doc, """           [SecureSync System]
    ┌───────────────────────────────────────────┐
    │                                           │
    │  (התחבר עם 2FA)      (צפה ב-Audit Logs)  │
    │  (העלה קובץ מוצפן)   (צפה בניקודי חריגה) │
    │  (הורד קובץ)  ←── ── (ניהול sessions)    │
    │  (שתף קובץ)          (חסום/שחרר מכשיר)   │
    │  (גרסה חדשה)         (חסום/שחרר משתמש)   │
    │  (מחק קובץ)                               │
    │                                           │
    └───────────────────────────────────────────┘
         ↑                         ↑
    [משתמש רגיל]              [מנהל (Admin)]""")

add_heading(doc, '15.8 רשימת Use Cases', level=2)
add_table(doc,
    ['UC', 'שם', 'Actor'],
    [
        ['UC-01', 'התחברות עם 2FA', 'משתמש'],
        ['UC-02', 'העלאת קובץ מוצפן', 'משתמש'],
        ['UC-03', 'הורדת קובץ ופענוח', 'משתמש / בעל הרשאה'],
        ['UC-04', 'שיתוף קובץ עם הרשאות', 'בעלים'],
        ['UC-05', 'העלאת גרסה חדשה', 'בעלים'],
        ['UC-06', 'מחיקת קובץ', 'בעלים'],
        ['UC-07', 'צפייה ב-Audit Logs', 'מנהל'],
        ['UC-08', 'צפייה בניקודי חריגה', 'מנהל'],
        ['UC-09', 'חסימת / שחרור מכשיר', 'מנהל'],
        ['UC-10', 'חסימת / שחרור משתמש', 'מנהל'],
        ['UC-11', 'צפייה בסשנים פעילים', 'מנהל'],
    ],
    col_widths=[1.5, 8, 4]
)

add_heading(doc, '15.9 תרשימי UML — Sequence Diagram: העלאת קובץ', level=2)
add_code(doc, """User        Client (React)     WebCrypto       Server         Python        MongoDB
 |                |                  |               |               |               |
 |--select file-->|                  |               |               |               |
 |                |--generateKey()-->|               |               |               |
 |                |<---AES_key-------|               |               |               |
 |                |---encrypt(file)->|               |               |               |
 |                |<-ciphertext+iv---|               |               |               |
 |                |---wrapKey()----->|               |               |               |
 |                |<-wrappedKey------|               |               |               |
 |                |---SHA256-------->|               |               |               |
 |                |<-hash------------|               |               |               |
 |                |----------POST /files/upload----->|               |               |
 |                |                                  |---verify SHA->|               |
 |                |                                  |---sign RSA--->|               |
 |                |                                  |----POST /analyze------------->|
 |                |                                  |<---{anomaly, zScore}----------|
 |                |                                  |--broadcast if anomaly         |
 |                |                                  |----------save File----------->|
 |                |<---------{id: fileId}------------|               |               |""")

add_heading(doc, '15.10 Design Class Diagram', level=2)
add_code(doc, """┌──────────────┐    1:N    ┌──────────────┐    1:N    ┌──────────────────┐
│    User      │──────────▶│    File      │──────────▶│   Permission     │
│──────────────│           │──────────────│           │──────────────────│
│ _id: ObjectId│           │ _id          │           │ fileId           │
│ email        │           │ userId (ref) │           │ grantedTo (ref)  │
│ passwordHash │           │ storedName   │           │ grantedBy (ref)  │
│ totpSecret   │           │ ciphertextSize│          │ role: enum       │
└──────┬───────┘           │ ciphertextSha│           │ expiresAt        │
       │ 1:N               │ algorithm    │           └──────────────────┘
       │                   │ ivB64        │
┌──────▼───────┐           │ wrappedKeyB64│    1:N    ┌──────────────────┐
│   Session    │           │ encryptedMeta│──────────▶│  FileVersion     │
│──────────────│           │ signature    │           │──────────────────│
│ userId (ref) │           └──────────────┘           │ fileId (ref)     │
│ jti: string  │                                       │ versionNumber    │
│ ip           │    1:N    ┌──────────────┐           │ storedName       │
└──────────────┘──────────▶│  AuditLog    │           └──────────────────┘
                           │──────────────│
┌──────────────┐           │ userId (ref) │    1:N    ┌──────────────────┐
│ AnomalyScore │           │ action       │──────────▶│    Device        │
│──────────────│           │ outcome      │           │──────────────────│
│ userId       │           │ ip           │           │ userId (ref)     │
│ zScore       │           │ detail       │           │ hostname, os     │
│ anomaly:bool │           └──────────────┘           │ openPorts[]      │
└──────────────┘                                       │ antivirus: bool  │
                                                       └──────────────────┘""")

add_heading(doc, '15.11 תרשים רכיבים', level=2)
add_code(doc, """┌─────────────────────────────────────────────────────────────┐
│                     Client Component                        │
│  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌──────────┐   │
│  │ App.jsx  │  │useApp.js │  │crypto.js │  │ api.js   │   │
│  │(UI state)│◄─│(state mgr)│  │(WebCrypto)│  │(HTTP)    │   │
│  └────┬─────┘  └──────────┘  └──────────┘  └────┬─────┘   │
│       │ WebSocket                                 │ HTTPS   │
└───────┼───────────────────────────────────────────┼─────────┘
        │ WSS                                       │ REST
┌───────▼───────────────────────────────────────────▼─────────┐
│                     Server Component                         │
│  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌──────────┐    │
│  │websocket │  │   auth   │  │  routes  │  │controllers│    │
│  │  .js     │  │middleware│  │ (6 files)│  │(6 files) │    │
│  └──────────┘  └──────────┘  └────┬─────┘  └────┬─────┘    │
│                                    │              │          │
│  ┌──────────┐  ┌──────────┐  ┌────▼─────┐  ┌────▼─────┐    │
│  │audit.js  │  │anomaly.js│  │  models  │  │  utils   │    │
│  └──────────┘  └────┬─────┘  │(20 files)│  │(pki.js)  │    │
└───────────────────┬─┼─────────┴──────────┴──────────────────┘
                    │ │ HTTP
┌───────────────────┴─▼──┐   ┌──────────────────────────────┐
│  Python Flask Service  │   │      MongoDB Atlas           │
│  Z-score, IF, SVM      │   │   20 collections             │
└────────────────────────┘   └──────────────────────────────┘""")

add_heading(doc, '15.12 תיאור הרכיבים (קומפוננטות) המוצעות', level=2)
add_rtl_para(doc, 'מסך כללי לפרטים — כל רכיב עיקרי:', bold=True)
add_table(doc,
    ['רכיב', 'תפקיד', 'קלטים', 'פלטים'],
    [
        ['App.jsx', 'ניהול state ראשי ו-routing', 'אירועי משתמש, WebSocket events', 'תצוגת מסכים'],
        ['crypto.js', 'כל פעולות הצפנה/פענוח', 'File object, RSA keys', 'ciphertext, ArrayBuffer'],
        ['useApp.js', 'ניהול state ו-API calls', 'אירועי UI', 'state updates'],
        ['useWebSocket.js', 'חיבור WebSocket + alerts', 'JWT token', 'anomaly alerts'],
        ['Dashboard.jsx', 'דשבורד ניהולי 5 טאבים', 'JWT, admin role', 'audit/anomaly/devices tables'],
        ['auth.js (middleware)', 'אימות JWT + BlockRule', 'HTTP request (Authorization header)', 'req.userId או 401/403'],
        ['audit.js', 'רישום פעולות', 'userId, action, outcome, ip, fileId', 'AuditLog doc in MongoDB'],
        ['anomaly.js', 'ניתוח חריגות', 'userId, action', 'AnomalyScore doc + WebSocket broadcast'],
        ['pki.js', 'חתימה ואימות RSA', 'buffer, private/public key', 'signature / boolean'],
    ],
    col_widths=[3.5, 4, 4.5, 3.5]
)
add_rtl_para(doc, 'זרימת מידע: משתמש → App.jsx → api.js → HTTP → auth middleware → controller → models → MongoDB')
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 16 — Screen Flow Diagram
# ══════════════════════════════════════════════════════════════
add_heading(doc, '16. תרשים מסכים (Screen Flow Diagram)')
add_code(doc, """[מסך Login / Register]
         │
         ├── Register → POST /users/register → OTP נשלח
         │
         └── Login → POST /users/login
                          │
                    [מסך OTP Verify]
                          │
                    POST /users/verify-otp
                          │
            ┌─────────────────────────────────┐
            │      [מסך קבצים ראשי]            │
            │                                  │
            ├── [Upload] → encrypt → POST      │
            ├── [Refresh] → GET /files/        │
            ├── [Download+Decrypt] → GET/:id   │
            │                  → decrypt → save│
            ├── [Share] → [Share Dialog]       │
            │    → POST /:id/share             │
            ├── [New Version] → [Version Dialog]│
            │    → POST /:id/versions          │
            ├── [Delete] → DELETE /:id         │
            └── [Dashboard Button]             │
                     │                         │
            [דשבורד ניהולי]                    │
            ├── Tab: Audit Logs                │
            ├── Tab: Anomaly Scores            │
            ├── Tab: Active Sessions           │
            ├── Tab: Devices → [Block/Unblock] │
            ├── Tab: Users  → [Block/Unblock]  │
            └── [← Back to Files]             │
            │                                  │
            └── [Logout] → POST /logout        │
                                               │
   [WebSocket Alert Banner] ←─────────────────┘
   (floating top-right, position: fixed)""")
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 17 — פירוט מסכים
# ══════════════════════════════════════════════════════════════
add_heading(doc, '17. פירוט מסכים')

add_heading(doc, '17.1 מסך התחברות / רישום', level=2)
add_rtl_para(doc, 'תפקיד: כניסה למערכת עם אמות אימות ראשוניות.')
add_rtl_para(doc, 'תיאור: מסך עם שני מצבים — login ו-register. בלחיצה על "Register" מוצגים שדות שם ואישור סיסמה. '
             'בלחיצה על "Login" נשלחת בקשת POST. הודעות שגיאה מוצגות מתחת לטופס.')
add_rtl_para(doc, 'צילום מסך: [יש לצרף צילום מסך של מסך הלוגין בעת הגשה]')

add_heading(doc, '17.2 מסך אימות OTP (2FA)', level=2)
add_rtl_para(doc, 'תפקיד: אימות זהות דו-שלבי.')
add_rtl_para(doc, 'תיאור: מסך עם שדה קוד 4 ספרות (inputMode="numeric"). מציג את כתובת האימייל שאליה נשלח הקוד. '
             'Auto-focus אוטומטי. כפתורי "אמת" ו-"חזרה".')
add_rtl_para(doc, 'צילום מסך: [יש לצרף צילום מסך של מסך OTP בעת הגשה]')

add_heading(doc, '17.3 מסך ניהול קבצים (ראשי)', level=2)
add_rtl_para(doc, 'תפקיד: ניהול כל הקבצים — העלאה, הורדה, שיתוף, גרסאות, מחיקה.')
add_rtl_para(doc, 'תיאור: מורכב מ-5 אזורים: (1) פס עליון עם כפתורי Dashboard ו-Logout, '
             '(2) אזור העלאה עם file input וכפתור "Encrypt + Upload", '
             '(3) banner אדום לאנומליות מ-WebSocket, '
             '(4) טבלת קבצים עם פעולות, '
             '(5) dialogs לשיתוף וגרסה חדשה.')
add_rtl_para(doc, 'צילום מסך: [יש לצרף צילום מסך של מסך הקבצים הראשי בעת הגשה]')

add_heading(doc, '17.4 דשבורד ניהולי', level=2)
add_rtl_para(doc, 'תפקיד: ניהול מערכת — אודיט, אנומליות, סשנים, מכשירים, משתמשים.')
add_rtl_para(doc, 'תיאור: 5 טאבים:')
add_rtl_para(doc, '• Audit Logs — User | Action | Result | IP | Time')
add_rtl_para(doc, '• Anomaly Scores — User | Action | Z-Score | Details | Time')
add_rtl_para(doc, '• Active Sessions — User | IP | Created | Expiry')
add_rtl_para(doc, '• Devices — Hostname | OS | IP | Ports | Antivirus | Status | Block')
add_rtl_para(doc, '• Users — Email | Name | Registered | Block/Unblock')
add_rtl_para(doc, 'צילום מסך: [יש לצרף צילום מסך של הדשבורד בעת הגשה]')
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 18 — אלמנטי תצוגה
# ══════════════════════════════════════════════════════════════
add_heading(doc, '18. אלמנטי תצוגה')
add_table(doc,
    ['אלמנט', 'מסך', 'תפקיד'],
    [
        ['כפתור "Encrypt + Upload"', 'קבצים ראשי', 'מפעיל את תהליך ההצפנה וההעלאה'],
        ['כפתור "Download+Decrypt"', 'קבצים ראשי', 'מוריד ciphertext ומפענח אוטומטית'],
        ['כפתור "Share"', 'קבצים ראשי', 'פותח dialog לשיתוף קובץ'],
        ['כפתור "New Version"', 'קבצים ראשי', 'פותח dialog להעלאת גרסה חדשה'],
        ['כפתור "Delete"', 'קבצים ראשי', 'מוחק קובץ לאחר אישור'],
        ['Banner אנומליה (אדום)', 'קבצים ראשי', 'מציג התראת אנומליה בזמן אמת מ-WebSocket'],
        ['כפתור "Dashboard"', 'קבצים ראשי', 'מנווט לדשבורד הניהולי'],
        ['כפתור "Logout"', 'כל המסכים', 'מבצע logout ומנקה JWT מ-localStorage'],
        ['שדה OTP (numeric input)', 'OTP Verify', 'קלט קוד 4 ספרות עם focus אוטומטי'],
        ['כפתורי טאב בדשבורד', 'דשבורד', 'מחליפים בין 5 views: Audit/Anomaly/Sessions/Devices/Users'],
        ['כפתור Block/Unblock', 'דשבורד', 'חוסם או משחרר מכשיר/משתמש (יוצר/מוחק BlockRule)'],
        ['role dropdown', 'Share Dialog', 'בחירת הרשאה: read / write / admin'],
    ],
    col_widths=[4.5, 3.5, 7.5]
)
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 19 — הודעות למשתמש
# ══════════════════════════════════════════════════════════════
add_heading(doc, '19. הודעות למשתמש')
add_table(doc,
    ['מצב', 'הודעה / Alert', 'סוג'],
    [
        ['OTP נשלח בהצלחה', '"OTP sent to your email"', 'הצלחה (ירוק)'],
        ['OTP שגוי', '"Invalid OTP"', 'שגיאה (אדום)'],
        ['JWT פג תוקף', '"Session expired. Please login again."', 'שגיאה'],
        ['משתמש חסום', '"User is blocked"', 'שגיאה 403'],
        ['JWT לאחר logout', '"Invalid or expired session"', 'שגיאה 401'],
        ['SHA-256 אינו תואם', '"File integrity check failed"', 'שגיאה 400'],
        ['RSA signature שגויה', '"File tampering detected"', 'שגיאה 403'],
        ['הצפנה בתהליך', '"Encrypting..."', 'Progress'],
        ['העלאה בתהליך', '"Uploading..."', 'Progress'],
        ['קובץ עלה בהצלחה', '"Uploaded: [filename]"', 'הצלחה'],
        ['אנומליה זוהתה (WebSocket)', 'Banner אדום: userId, action, Z-score', 'Alert (WebSocket)'],
        ['ניסיון גישה ללא הרשאה', '"Access denied"', 'שגיאה 403'],
    ],
    col_widths=[5, 7, 3.5]
)
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 20 — ממשק משתמש
# ══════════════════════════════════════════════════════════════
add_heading(doc, '20. ממשק משתמש')
add_rtl_para(doc, 'ממשק המשתמש בנוי ב-React 18 עם CSS ישיר (ללא UI framework חיצוני).')
add_rtl_para(doc, '• שפה: ממשק באנגלית (טבלאות, כפתורים) עם תמיכה בעברית בשמות קבצים')
add_rtl_para(doc, '• Layout: SPA (Single Page Application) — מסך אחד שמשנה את ה-state')
add_rtl_para(doc, '• עיצוב: רקע כהה (#1a1a2e), כרטיסים בהירים, כפתורים בצבע #4a90d9')
add_rtl_para(doc, '• Responsive: עמודות מסתגלות לגודל המסך')
add_rtl_para(doc, '• נגישות: כל שדות הקלט עם placeholder ו-label ברור')
add_rtl_para(doc, '• WebSocket Alert Banner: position: fixed, top-right, z-index: 9999, צבע אדום #c62828')
add_rtl_para(doc, '• State Machine: App מנהל 3 states — login | verify | files')
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 21 — קוד התוכנית
# ══════════════════════════════════════════════════════════════
add_heading(doc, '21. קוד התוכנית עם תיעוד')

add_heading(doc, '21.1 קלטים עיקריים', level=2)
add_code(doc, """// קלטים מהמשתמש
POST /users/register:  { email: string, name: string, password: string }
POST /users/login:     { email: string, password: string }
POST /users/verify-otp:{ email: string, otp: string }

// קלט קובץ (multipart/form-data)
POST /files/upload:
  file:               binary (ciphertext)
  ivB64:              string (base64, 12 bytes IV)
  wrappedKeyB64:      string (base64, RSA-wrapped AES key)
  encryptedMetaB64:   string (base64, encrypted metadata)
  metaIvB64:          string (base64, metadata IV)
  ciphertextSha256B64:string (base64, SHA-256 hash)
  algorithm:          "AES-GCM"

// קלט שיתוף
POST /files/:id/share: { email: string, role: "read"|"write"|"admin" }

// קלט Python /analyze
POST /analyze: { userId: string, action: string }""")

add_heading(doc, '21.2 פלטים עיקריים', level=2)
add_code(doc, """// תגובות שרת
POST /users/verify-otp → { token: "eyJ..." }   // JWT
GET  /files/           → [ { _id, storedName, encryptedMetaB64,
                             wrappedKeyB64, ivB64, algorithm,
                             metaIvB64, createdAt, isOwner } ]
GET  /files/:id/download →
  body:   binary ciphertext
  header: x-securesync-meta: base64(JSON({wrappedKeyB64, encryptedMetaB64, metaIvB64}))

// תגובה Python /analyze
{ "anomaly": bool, "zScore": float, "ewma": float,
  "details": string, "score": float }""")

add_heading(doc, '21.3 פונקציות חשובות / קריטיות', level=2)
add_code(doc, """// crypto.js — encryptFileWithWrappedKey(file)
// הצפנת קובץ ו-metadata, עטיפת מפתח AES עם RSA-OAEP
async function encryptFileWithWrappedKey(file) {
  const aesKey = await crypto.subtle.generateKey({name:"AES-GCM",length:256},true,["encrypt","decrypt"]);
  const iv = crypto.getRandomValues(new Uint8Array(12));
  const ciphertext = await crypto.subtle.encrypt({name:"AES-GCM",iv},aesKey,await file.arrayBuffer());
  const metaJson = JSON.stringify({name:file.name, type:file.type, size:file.size});
  const metaIv = crypto.getRandomValues(new Uint8Array(12));
  const encMeta = await crypto.subtle.encrypt({name:"AES-GCM",iv:metaIv},aesKey,
                    new TextEncoder().encode(metaJson));
  const rsaKeyPair = await getOrCreateRsaKeyPair();
  const wrappedKey = await crypto.subtle.wrapKey("raw",aesKey,rsaKeyPair.publicKey,"RSA-OAEP");
  const hash = await crypto.subtle.digest("SHA-256",ciphertext);
  return { ciphertext, iv, wrappedKey, encMeta, metaIv, hash };
}

// auth.js — authRequired middleware
// אימות JWT + בדיקת jti במסד נתונים + בדיקת BlockRule
async function authRequired(req, res, next) {
  const token = req.headers.authorization?.split(' ')[1];
  const decoded = jwt.verify(token, process.env.JWT_SECRET);
  const session = await Session.findOne({ jti: decoded.jti });
  if (!session) return res.status(401).json({error: "Invalid or expired session"});
  const blocked = await BlockRule.findOne({userId: decoded.userId, active: true});
  if (blocked) return res.status(403).json({error: "User is blocked"});
  req.userId = decoded.userId;
  next();
}

// anomaly.js — analyzeEvent(userId, action)
// שולח לPython service ושמר תוצאה + broadcast אם anomaly
async function analyzeEvent(userId, action) {
  const res = await fetch('http://localhost:5001/analyze', {
    method:'POST', headers:{'Content-Type':'application/json'},
    body: JSON.stringify({userId, action})
  });
  const result = await res.json();
  await AnomalyScore.create({userId, action, zScore:result.zScore, anomaly:result.anomaly,...});
  if (result.anomaly) broadcast({type:'anomaly', userId, action, ...result});
}""")
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 22 — מסד הנתונים
# ══════════════════════════════════════════════════════════════
add_heading(doc, '22. תיאור מסד הנתונים')
add_rtl_para(doc, 'מסד הנתונים: MongoDB Atlas — NoSQL document store. 20 Mongoose schemas.')
add_rtl_para(doc, 'ERD — יחסים בין ישויות:', bold=True)
add_code(doc, """User ──< Session    (1:N) — משתמש יכול לפתוח מספר sessions
User ──< File        (1:N) — משתמש בעלים של קבצים
User ──< AuditLog   (1:N) — כל פעולה מתועדת
User ──< Device      (1:N) — מכשירים של המשתמש
User ──< KeyPair     (1:1) — זוג מפתחות RSA לכל משתמש
File ──< Permission  (1:N) — קובץ משותף למספר משתמשים
File ──< FileVersion (1:N) — גרסאות של קובץ""")

add_rtl_para(doc, 'פירוט טבלאות (Collections):', bold=True)
add_table(doc,
    ['אוסף (Collection)', 'עמודות עיקריות', 'מפתח ראשי / זר', 'סוג נתונים', 'חובה'],
    [
        ['users', '_id, email, name, passwordHash, totpSecret', '_id (ObjectId)', 'String, String, String', 'email ✓'],
        ['files', '_id, userId, storedName, ciphertextSize, algorithm, ivB64, wrappedKeyB64, encryptedMetaB64, ciphertextSha256B64, signature', 'userId → users', 'ref, String, Number, String', 'userId, storedName ✓'],
        ['permissions', '_id, fileId, grantedTo, grantedBy, role, expiresAt, revokedAt', 'fileId→files, grantedTo→users', 'ref, enum(read/write/admin)', 'fileId, grantedTo, role ✓'],
        ['auditlogs', '_id, userId, action, outcome, ip, fileId, detail', 'userId → users', 'ref, String, String, String', 'userId, action, outcome ✓'],
        ['sessions', '_id, userId, jti, ip, expiresAt', 'userId → users', 'ref, String (unique), Date', 'userId, jti ✓'],
        ['anomalyscores', '_id, userId, action, zScore, ewma, anomaly, details', 'userId (ref)', 'ref, String, Float, Boolean', 'userId, action ✓'],
        ['devices', '_id, userId, hostname, os, ipAddress, openPorts, antivirus, status, lastSeen', 'userId → users', 'ref, String[], Boolean, enum(online/offline)', 'userId, hostname ✓'],
        ['blockrules', '_id, userId, hostname, ipAddress, reason, blockedBy, active', 'userId→users, blockedBy→users', 'ref, String, Boolean', 'active ✓'],
        ['fileversions', '_id, fileId, versionNumber, storedName, ciphertextSize, algorithm, ivB64, wrappedKeyB64', 'fileId → files', 'ref, Number, String', 'fileId, versionNumber ✓'],
        ['keypairs', '_id, userId, publicKey, privateKeyEnc', 'userId → users (unique)', 'ref, String (PEM)', 'userId, publicKey ✓'],
    ],
    col_widths=[3.5, 5, 3, 3, 1.5]
)
add_rtl_para(doc, 'כל collection משתמש בתצורת timestamps: true (createdAt, updatedAt אוטומטיים). '
             'לא נעשה שימוש ב-stored procedures או views (MongoDB לא תומך בכך). '
             'שדות עם index: userId, jti (unique), email (unique).')
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 23 — מדריך למשתמש
# ══════════════════════════════════════════════════════════════
add_heading(doc, '23. מדריך למשתמש')
add_rtl_para(doc, 'הפעלת המערכת:', bold=True)
add_code(doc, """# 1. הפעלת שרת Node.js
cd C:/fullstack/securesync/Server
npm run dev   (→ HTTPS on port 4000)

# 2. הפעלת Python ML service
cd Server/anomaly_service
python main.py   (→ Flask on port 5001)

# 3. הפעלת קליינט React
cd C:/fullstack/securesync/Client
npm run dev   (→ Vite on port 5173)

# 4. גישה דרך הדפדפן
https://localhost:5173
(בפעם הראשונה: לחץ "Advanced" ← "thisisunsafe" בשורת הכתובת)""")

add_rtl_para(doc, 'שימוש במערכת:', bold=True)
add_rtl_para(doc, '1. רישום: הזן email, שם, סיסמה → לחץ "Register"')
add_rtl_para(doc, '2. OTP: בדוק אימייל → הזן קוד 4 ספרות → לחץ "Verify"')
add_rtl_para(doc, '3. העלאת קובץ: לחץ "Choose File" → בחר קובץ → לחץ "Encrypt + Upload"')
add_rtl_para(doc, '4. הורדת קובץ: לחץ "Download+Decrypt" ליד הקובץ → הקובץ המפוענח מורד')
add_rtl_para(doc, '5. שיתוף: לחץ "Share" → הזן email + role → לחץ "Share File"')
add_rtl_para(doc, '6. גרסה חדשה: לחץ "New Version" → בחר קובץ חדש → לחץ "Upload"')
add_rtl_para(doc, '7. דשבורד: לחץ "Dashboard" → ניווט בין הטאבים')
add_rtl_para(doc, '8. יציאה: לחץ "Logout" — ה-session נמחק מהשרת')
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 24 — בדיקות
# ══════════════════════════════════════════════════════════════
add_heading(doc, '24. בדיקות והערכה')

add_heading(doc, '24.1 בדיקות פונקציונליות', level=2)
add_table(doc,
    ['#', 'תרחיש', 'פעולה', 'תוצאה צפויה', 'תוצאה בפועל'],
    [
        ['1', 'רישום משתמש', 'POST /users/register', 'OTP נשלח לאימייל', '✓'],
        ['2', 'OTP שגוי', 'POST /users/verify-otp', 'שגיאה 400', '✓'],
        ['3', 'OTP נכון', 'POST /users/verify-otp', 'JWT מתקבל', '✓'],
        ['4', 'העלאת קובץ', 'POST /files/upload', '{id: fileId}', '✓'],
        ['5', 'רשימת קבצים', 'GET /files/', 'מערך קבצים', '✓'],
        ['6', 'הורדת קובץ', 'GET /files/:id/download', 'ciphertext + header', '✓'],
        ['7', 'פענוח קובץ', 'WebCrypto client-side', 'קובץ מקורי', '✓'],
        ['8', 'שיתוף קובץ', 'POST /files/:id/share', 'Permission נוצר', '✓'],
        ['9', 'גישה לקובץ משותף', 'GET /files/ (user B)', 'קובץ מופיע', '✓'],
        ['10', 'גרסה חדשה', 'POST /files/:id/versions', 'גרסה נשמרת', '✓'],
        ['11', 'מחיקת קובץ', 'DELETE /files/:id', 'נמחק מ-DB ומדיסק', '✓'],
        ['12', 'Logout', 'POST /users/logout', 'Session נמחק', '✓'],
    ],
    col_widths=[1, 4, 4, 4.5, 2]
)

add_heading(doc, '24.2 בדיקות אבטחה', level=2)
add_table(doc,
    ['#', 'תרחיש', 'תוצאה צפויה', 'תוצאה בפועל'],
    [
        ['1', 'גישה לקובץ של משתמש אחר', '403 Access denied', '✓'],
        ['2', 'JWT פג תוקף', '401 Unauthorized', '✓'],
        ['3', 'JWT אחרי logout', '401 Invalid session', '✓'],
        ['4', 'SHA-256 לא תואם', '400 Hash mismatch', '✓'],
        ['5', 'RSA signature שגויה', '403 Tampering detected', '✓'],
        ['6', 'משתמש חסום', '403 User is blocked', '✓'],
        ['7', 'מכשיר חסום', '403 Device is blocked', '✓'],
        ['8', 'בקשה ללא JWT', '401 No token', '✓'],
    ],
    col_widths=[1, 6, 4.5, 2]
)

add_heading(doc, '24.3 בדיקת Anomaly Detection', level=2)
add_rtl_para(doc, 'שלבי הבדיקה:')
add_rtl_para(doc, '1. הכנסת 59 לוגי AuditLog היסטוריים כ-baseline')
add_rtl_para(doc, '2. ביצוע פעולות download רבות בזמן קצר')
add_rtl_para(doc, '3. Python service חישב Z-score = 3.44 (מעל סף 2.0)')
add_rtl_para(doc, '4. anomaly = true → WebSocket broadcast')
add_rtl_para(doc, '5. Banner אדום הופיע בקליינט בזמן אמת ✓')
add_rtl_para(doc, '6. AnomalyScore נשמר ב-DB ומוצג בדשבורד ✓')
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 25 — מסקנות
# ══════════════════════════════════════════════════════════════
add_heading(doc, '25. מסקנות')
add_rtl_para(doc, 'SecureSync מממש מערכת אבטחת קבצים ארגונית מלאה עם 7 שכבות הגנה (Defense in Depth). '
             'כל 9 מטרות הפרויקט הושגו.')
add_rtl_para(doc, 'לקחים עיקריים:', bold=True)
add_rtl_para(doc, '• WebCrypto API עוצמתי — הצפנה חזקה בדפדפן ללא ספריות חיצוניות')
add_rtl_para(doc, '• Zero-Knowledge דורש תכנון קפדני — כל metadata חייב להיות מוצפן בנפרד')
add_rtl_para(doc, '• Anomaly Detection דורש tuning — ספי גילוי ונחוצים לאיזון בין False Positives ל-False Negatives')
add_rtl_para(doc, '• Defense in Depth — אמינות מגיעה מריבוי שכבות עצמאיות, לא שכבה אחת מושלמת')
add_rtl_para(doc, '• Node.js async I/O מתאים מאוד לשרת קבצים — תפוקה גבוהה עם concurrency')
add_rtl_para(doc, '• MongoDB גמיש לפיתוח מהיר — 20 models בלי schema migrations')
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 26 — פיתוחים עתידיים
# ══════════════════════════════════════════════════════════════
add_heading(doc, '26. פיתוחים עתידיים')
add_table(doc,
    ['פיתוח', 'תיאור', 'עדיפות'],
    [
        ['PKI מלא', 'Certificate Authority פנימי + mTLS בין Agent לשרת', 'גבוהה'],
        ['Mobile App', 'React Native app עם WebCrypto polyfill', 'גבוהה'],
        ['Deployment לענן', 'AWS/Azure deployment עם TLS אמיתי (Let\'s Encrypt)', 'גבוהה'],
        ['Password reset', 'זרימת שחזור סיסמה מאובטחת', 'בינונית'],
        ['File size quota', 'מגבלת גודל ו-quota למשתמש', 'בינונית'],
        ['Endpoint Agent', 'תוכנת agent ב-OS לניטור אנטיוירוס, פורטים', 'נמוכה'],
        ['Real-time router blocking', 'חסימת מכשיר ב-switch/router רשת', 'נמוכה'],
        ['ML model persistence', 'שמירת מודלי Isolation Forest/SVM לאחר אימון', 'נמוכה'],
    ],
    col_widths=[4, 9, 2.5]
)
page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 27 — ביבליוגרפיה
# ══════════════════════════════════════════════════════════════
add_heading(doc, '27. ביבליוגרפיה')
refs = [
    'MDN Web Docs — Web Crypto API. https://developer.mozilla.org/en-US/docs/Web/API/Web_Crypto_API',
    'NIST SP 800-38D — Recommendation for Block Cipher Modes of Operation: GCM. 2007.',
    'RFC 7519 — JSON Web Tokens (JWT). IETF. 2015.',
    'OWASP Top 10. https://owasp.org/www-project-top-ten/',
    'Chandola, V., Banerjee, A., & Kumar, V. (2009). Anomaly detection: A survey. ACM Computing Surveys.',
    'Liu, F. T., Ting, K. M., & Zhou, Z. H. (2008). Isolation Forest. IEEE ICDM.',
    'Schölkopf, B. et al. (2001). Estimating the Support of a High-Dimensional Distribution. Neural Computation.',
    'MongoDB Mongoose Documentation. https://mongoosejs.com/docs/',
    'scikit-learn Documentation. https://scikit-learn.org/stable/modules/outlier_detection.html',
    'Express.js Documentation. https://expressjs.com/',
    'React Documentation. https://react.dev/',
    'Vite Build Tool Documentation. https://vitejs.dev/',
    'ws — WebSocket Node.js. https://github.com/websockets/ws',
    'jsonwebtoken npm. https://github.com/auth0/node-jsonwebtoken',
    'bcrypt npm. https://github.com/kelektiv/node.bcrypt.js',
    'Nodemailer Documentation. https://nodemailer.com/',
]
for i, ref in enumerate(refs, 1):
    add_rtl_para(doc, f'{i}. {ref}', size=11)

# Save
output_path = r'C:/fullstack/securesync/ספר_פרויקט_SecureSync.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
